"""Zotero 活动引用兼容层（src/docx_io/zotero_fields.py）的回归测试。"""

from __future__ import annotations

import json
import re
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from xml.sax.saxutils import escape

from src.docx_io.zotero_fields import (
    docx_has_zotero_fields,
    repair_docx,
    strip_numbering_inside_zotero_bibliography,
)

CONTENT_TYPES = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
    '<Default Extension="xml" ContentType="application/xml"/>'
    '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
    "</Types>"
)
PACKAGE_RELS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
    "</Relationships>"
)
DOCUMENT_RELS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rIdZoteroData" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" Target="../customXml/item2.xml"/>'
    "</Relationships>"
)

# 数字 id、缺失 uris、fldSimple —— 三种 Zotero 9 会直接报错的写法
BROKEN_PAYLOAD = {
    "citationID": "cTest01",
    "properties": {"formattedCitation": "[1]", "plainCitation": "[1]", "style": "x"},
    "citationItems": [
        {"id": 1, "itemData": {"id": 1, "type": "book", "title": "国史旧闻"}}
    ],
    "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
}


def _broken_document_xml() -> str:
    instr = "ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(
        BROKEN_PAYLOAD, ensure_ascii=False, separators=(",", ":")
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        f'<w:p><w:fldSimple w:instr="{escape(instr, {chr(34): "&quot;"})}">'
        "<w:r><w:t>[1]</w:t></w:r></w:fldSimple></w:p>"
        "<w:p><w:r><w:t>ADDIN ZOTERO_PREF stale</w:t></w:r></w:p>"
        "<w:p><w:r><w:t>参考文献</w:t></w:r></w:p>"
        "</w:body></w:document>"
    )


def _make_broken_docx(path: Path) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", CONTENT_TYPES)
        archive.writestr("_rels/.rels", PACKAGE_RELS)
        archive.writestr("word/_rels/document.xml.rels", DOCUMENT_RELS)
        archive.writestr("word/document.xml", _broken_document_xml())
        archive.writestr("customXml/item2.xml", "<data>ZOTERO_PREF junk</data>")
        archive.writestr("customXml/_rels/item2.xml.rels", "<Relationships/>")


def _read(path: Path, name: str) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read(name).decode("utf-8")


class ZoteroCompatTests(unittest.TestCase):
    def test_repair_makes_document_refreshable(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "broken.docx"
            _make_broken_docx(path)
            self.assertTrue(docx_has_zotero_fields(path))

            report = repair_docx(path)
            self.assertTrue(report.changed)
            self.assertEqual(report.citations, 1)
            self.assertEqual(report.field_simple_converted, 1)
            self.assertTrue(report.prefs_written)

            xml = _read(path, "word/document.xml")
            # fldSimple → 复杂域
            self.assertNotIn("fldSimple", xml)
            self.assertIn('w:fldCharType="begin"', xml)
            self.assertIn('w:fldCharType="separate"', xml)
            self.assertIn('w:fldCharType="end"', xml)
            # 正文里的 ZOTERO_PREF 段落被清掉
            self.assertNotIn("ZOTERO_PREF", xml)

            payload = json.loads(
                re.search(r"CSL_CITATION\s+(\{.*?\})\s*</w:instrText>", xml, re.DOTALL).group(1)
            )
            item = payload["citationItems"][0]
            self.assertEqual(item["uris"], [])
            self.assertEqual(item["id"], "1")
            self.assertEqual(item["itemData"]["id"], "1")
            self.assertNotIn("style", payload["properties"])
            self.assertEqual(payload["properties"]["noteIndex"], 0)

            # 文档首选项写进 Word 自定义属性
            custom = _read(path, "docProps/custom.xml")
            self.assertIn("ZOTERO_PREF_1", custom)
            self.assertIn("china-national-standard-gb-t-7714-2015-numeric", custom)
            self.assertIn(
                'PartName="/docProps/custom.xml"', _read(path, "[Content_Types].xml")
            )
            self.assertIn("docProps/custom.xml", _read(path, "_rels/.rels"))
            self.assertNotIn("item2.xml", _read(path, "word/_rels/document.xml.rels"))

            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
            self.assertNotIn("customXml/item2.xml", names)

    def test_repair_is_idempotent(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "broken.docx"
            _make_broken_docx(path)
            repair_docx(path)
            first = path.read_bytes()
            second_report = repair_docx(path)
            self.assertFalse(second_report.changed)
            self.assertEqual(first, path.read_bytes())

    def test_documents_without_zotero_fields_are_untouched(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "plain.docx"
            with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
                archive.writestr("[Content_Types].xml", CONTENT_TYPES)
                archive.writestr("_rels/.rels", PACKAGE_RELS)
                archive.writestr(
                    "word/document.xml",
                    '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    "<w:body><w:p><w:r><w:t>正文</w:t></w:r></w:p></w:body></w:document>",
                )
            before = path.read_bytes()
            report = repair_docx(path)
            self.assertFalse(report.has_fields)
            self.assertFalse(report.changed)
            self.assertEqual(before, path.read_bytes())

    def test_default_scene_enables_zotero_compat(self) -> None:
        from src.scene.manager import load_scene

        preset = Path(__file__).resolve().parents[1] / "src/scene/presets/default_format.json"
        config = load_scene(preset)
        self.assertTrue(config.zotero.enabled)
        self.assertTrue(config.capabilities.get("zotero_live_citation"))
        self.assertIn("gb-t-7714-2015-numeric", config.zotero.style_id)

    def test_builtin_resources_are_present(self) -> None:
        from src.resources import zotero_assets

        assets = zotero_assets()
        self.assertIn("template", assets)
        self.assertIn("csl", assets)
        self.assertTrue(assets["csl"].read_text(encoding="utf-8").lstrip().startswith("<?xml"))
        self.assertTrue(docx_has_zotero_fields(assets["template"]))


def _seq_field(number: str) -> str:
    return (
        '<w:r><w:t xml:space="preserve">[</w:t></w:r>'
        f'<w:bookmarkStart w:id="1" w:name="_RefNum_RefEntry_{number}"/>'
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> SEQ RefEntry \\* ARABIC </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        f'<w:r><w:t>{number}</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        f'<w:bookmarkEnd w:id="1"/>'
        '<w:r><w:t xml:space="preserve">] </w:t></w:r>'
    )


BIBL_WITH_SEQ = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    "<w:body>"
    # 域首段：SEQ 编号被挤到 Zotero 域代码之前（Word 里就会显示成 [1] [1]）
    "<w:p>"
    + _seq_field("1")
    + '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
    '<w:r><w:instrText xml:space="preserve"> ADDIN ZOTERO_BIBL {"uncited":[]} CSL_BIBLIOGRAPHY </w:instrText></w:r>'
    '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
    "<w:r><w:t>陈登原. 国史旧闻.</w:t></w:r>"
    "</w:p>"
    # 域内后续段落：SEQ 编号插在域内容里
    "<w:p>" + _seq_field("2") + "<w:r><w:t>袁训来. 蓝田生物群.</w:t></w:r>"
    '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
    # 域外的普通参考文献段落：不能被动到
    "<w:p>" + _seq_field("3") + "<w:r><w:t>附录条目.</w:t></w:r></w:p>"
    "</w:body></w:document>"
)


class ZoteroBibliographyNumberingTests(unittest.TestCase):
    def test_seq_numbering_moved_into_bibliography_field(self) -> None:
        fixed, removed = strip_numbering_inside_zotero_bibliography(BIBL_WITH_SEQ)
        self.assertEqual(removed, 2)

        # 域内的两条编号被还原成静态文本，且都位于域内容里
        bibl_start = fixed.index("ZOTERO_BIBL")
        para_start = fixed.rfind("<w:p>", 0, bibl_start)
        before_field = fixed[para_start:bibl_start]
        self.assertNotIn("SEQ RefEntry", before_field)
        self.assertNotIn("[1]", before_field)

        separate = fixed.index('w:fldCharType="separate"', bibl_start)
        after_separate = fixed[separate : separate + 400]
        self.assertIn("[1] ", after_separate)

        # 域外那条参考文献的 SEQ 编号保持不变
        tail = fixed[fixed.index("附录条目") - 800 : ]
        self.assertIn("SEQ RefEntry", tail)
        self.assertEqual(fixed.count("SEQ RefEntry"), 1)

    def test_stripper_is_noop_without_zotero_bibliography(self) -> None:
        plain = BIBL_WITH_SEQ.replace("ZOTERO_BIBL", "OTHER_FIELD")
        fixed, removed = strip_numbering_inside_zotero_bibliography(plain)
        self.assertEqual(removed, 0)
        self.assertEqual(fixed, plain)

    def test_citation_link_rule_locks_zotero_paragraphs(self) -> None:
        from docx import Document as DocxDocument
        from lxml import etree

        from src.engine.rules.citation_link import _zotero_locked_paragraph_indexes

        doc = DocxDocument()
        doc.add_paragraph("正文段落")
        para = doc.add_paragraph()
        para._element.append(
            etree.fromstring(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:fldChar w:fldCharType="begin"/></w:r>'
            )
        )
        para._element.append(
            etree.fromstring(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:instrText xml:space="preserve"> ADDIN ZOTERO_BIBL {} CSL_BIBLIOGRAPHY </w:instrText>'
                "</w:r>"
            )
        )
        second = doc.add_paragraph("域内第二条")
        second._element.append(
            etree.fromstring(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:fldChar w:fldCharType="end"/></w:r>'
            )
        )
        doc.add_paragraph("域外段落")

        locked = _zotero_locked_paragraph_indexes(doc)
        self.assertIn(1, locked)
        self.assertIn(2, locked)
        self.assertNotIn(0, locked)
        self.assertNotIn(3, locked)


if __name__ == "__main__":
    unittest.main()
